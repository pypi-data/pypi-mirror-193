import unittest
from ddt import ddt, data
from docx import Document
from digestparser import parse
from digestparser.parse import LINE_SEPARATOR
from tests import read_fixture, data_path


@ddt
class TestParse(unittest.TestCase):
    def setUp(self):
        pass

    def test_parse_content(self):
        "test parsing all the content"
        expected = read_fixture("digest_99999.txt").decode("utf-8")
        content = parse.parse_content(data_path("DIGEST 99999.docx"))
        self.assertEqual(content, expected)

    def test_html_open_close_tag_failure(self):
        "test getting an open and close tag for a style that does not exist"
        style = "not_a_style"
        expected = None, None
        self.assertEqual(parse.html_open_close_tag(style), expected)

    def test_open_close_style_failure(self):
        "test wrapping content in close and open tags for a style that does not exist"
        run = None
        prev_run = None
        output = ""
        style = "not_a_style"
        expected = output
        self.assertEqual(
            parse.run_open_close_style(run, prev_run, output, style), expected
        )

    @data(
        {
            "scenario": "",
            "one_has_attr": None,
            "two_has_attr": None,
            "one_contains_break": None,
            "output": "",
            "attribute": "italic",
            "expected": "",
        },
        {
            "scenario": "previous run is italic, current run is not",
            "one_has_attr": True,
            "two_has_attr": None,
            "one_contains_break": None,
            "output": "",
            "attribute": "italic",
            "expected": "</i>",
        },
        {
            "scenario": "previous run is italic, current run is also italic",
            "one_has_attr": True,
            "two_has_attr": True,
            "one_contains_break": None,
            "output": "",
            "attribute": "italic",
            "expected": "",
        },
        {
            "scenario": "previous run is not italic, current run is italic",
            "one_has_attr": None,
            "two_has_attr": True,
            "one_contains_break": None,
            "output": "",
            "attribute": "italic",
            "expected": "<i>",
        },
        {
            "scenario": (
                "previous run is not italic, current run is italic, "
                + "output ends in a new line"
            ),
            "one_has_attr": None,
            "two_has_attr": True,
            "one_contains_break": None,
            "output": "test\n",
            "attribute": "italic",
            "expected": "test\n<i>",
        },
        {
            "scenario": (
                "previous run italic and contains a break, current run italic, "
                + "output ends in a new line"
            ),
            "one_has_attr": True,
            "two_has_attr": True,
            "one_contains_break": True,
            "output": "test\n",
            "attribute": "italic",
            "expected": "test</i>\n<i>",
        },
    )
    def test_open_close_style(self, test_data):
        "test the open close tag logic with basic attributes"
        self.assertEqual(
            parse.open_close_style(
                test_data.get("one_has_attr"),
                test_data.get("two_has_attr"),
                test_data.get("one_contains_break"),
                test_data.get("output"),
                test_data.get("attribute"),
            ),
            test_data.get("expected"),
        )

    def test_join_runs_bold_italic(self):
        "test to join bold run before an italic run"
        document = Document()
        paragraph = document.add_paragraph("")
        paragraph.add_run("bold ").bold = True
        paragraph.add_run("italic.").italic = True
        output = parse.join_runs(paragraph.runs)
        self.assertEqual(output, "<b>bold</b> <i>italic.</i>")

    def test_join_runs_italic_bold(self):
        "test to join italic run before a bold run"
        document = Document()
        paragraph = document.add_paragraph("")
        paragraph.add_run("italic ").italic = True
        paragraph.add_run("bold.").bold = True
        output = parse.join_runs(paragraph.runs)
        self.assertEqual(output, "<i>italic</i> <b>bold.</b>")

    def test_join_blank_bold_tag(self):
        """test joining runs with last run is a bold space"""
        document = Document()
        paragraph = document.add_paragraph("")
        paragraph.add_run("Author Name")
        paragraph.add_run(" ").bold = True
        output = parse.join_runs(paragraph.runs)
        # space will be retained but not the bold tag
        self.assertEqual(output, "Author Name ")

    def test_odd_whitespace(self):
        """test joining runs with strange whitespace based on a real example"""
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("AUTHOR").bold = True
        paragraph.add_run("\n")
        paragraph.add_run("Author Name")
        paragraph2 = document.add_paragraph()
        paragraph2.add_run("DIGEST").bold = True
        paragraph2.add_run(" ").bold = True
        paragraph2.add_run("TITLE").bold = True
        paragraph2.add_run("\n").bold = True
        # \u2028 will be stripped out
        paragraph2.add_run(LINE_SEPARATOR).bold = True
        paragraph2.add_run("\u201C")
        paragraph2.add_run("A title")
        paragraph2.add_run("\u201D")
        output = parse.parse_paragraphs(document)
        # space will be retained but not the bold tag
        self.assertEqual(
            output,
            "<b>AUTHOR</b>\nAuthor Name\n<b>DIGEST TITLE</b>\n\u201CA title\u201D\n",
        )


if __name__ == "__main__":
    unittest.main()
