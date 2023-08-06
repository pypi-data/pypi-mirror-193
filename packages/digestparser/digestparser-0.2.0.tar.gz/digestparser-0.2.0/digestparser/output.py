"build DOCX output from digest content"

from bs4 import BeautifulSoup
from docx import Document
from digestparser import utils
from digestparser.build import build_digest


class RunPart:
    "A run part to hold text and style attributes"

    def __new__(cls):
        new_instance = object.__new__(cls)
        new_instance.__init__()
        return new_instance

    def __init__(self):
        self.text = None
        self.italic = None
        self.bold = None
        self.subscript = None
        self.superscript = None


def html_runs(html):
    "convert HTML to text and run style attributes from very clean input HTML"
    run_parts = []
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup:
        part = RunPart()
        if not tag.name:
            part.text = tag
            run_parts.append(part)
        else:
            # part = RunPart()
            part.text = tag.text
            if tag.name == "i":
                part.italic = True
            if tag.name == "b":
                part.bold = True
            if tag.name == "sub":
                part.subscript = True
            if tag.name == "sup":
                part.superscript = True
            run_parts.append(part)
    return run_parts


def docx_file_name(digest, digest_config=None):
    "name for the docx output file formatted from the config if available"
    default_file_name_pattern = "{author}_{msid:0>5}.docx"
    file_name_pattern = default_file_name_pattern
    if digest_config and "output_file_name_pattern" in digest_config:
        file_name_pattern = str(digest_config.get("output_file_name_pattern"))
    # collect the values from the digest if present
    file_name = file_name_pattern.format(
        author=digest.author, msid=str(utils.msid_from_doi(digest.doi))
    )
    return utils.sanitise_file_name(file_name)


def digest_docx(digest, output_file_name):
    "convert a digest object to DOCX output"
    # create the docx Document
    document = Document()
    # add a header line
    paragraph = document.add_paragraph("DIGEST")
    # add the paragraphs of text
    for para in digest.text:
        paragraph = document.add_paragraph()
        run_parts = html_runs(para)
        for run_part in run_parts:
            run = paragraph.add_run(run_part.text)
            font = run.font
            if run_part.italic:
                font.italic = True
            if run_part.bold:
                font.bold = True
            if run_part.subscript:
                font.subscript = True
            if run_part.superscript:
                font.superscript = True
    # save the file
    with open(output_file_name, "wb") as open_file:
        document.save(open_file)
    return output_file_name


def build_docx(file_name, output_file_name):
    "build an output DOCX from a DOCX input file"
    digest = build_digest(file_name)
    docx_file = digest_docx(digest, output_file_name)
    return docx_file
